"""Render a markdown file to a styled .docx using python-docx.
Supports: H1-H4, paragraphs, bullet/numbered lists, fenced code blocks,
inline `code`, **bold**, *italic*, GFM pipe tables.
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1] if len(sys.argv) > 1 else r'D:/Projects/RWR/mvp/.qa-asterra-doc.md'
DST = sys.argv[2] if len(sys.argv) > 2 else r'D:/Projects/RWR/ASTERA_Documentation.updated.docx'

doc = Document()
# Tighten default styles a touch — this is a reference doc, dense layout.
for s in doc.styles:
    try:
        if s.name in ('Normal', 'List Bullet', 'List Number'):
            s.font.name = 'Calibri'
            s.font.size = Pt(10)
    except Exception:
        pass
# Heading sizes
for n, sz in (('Heading 1', 20), ('Heading 2', 16), ('Heading 3', 13), ('Heading 4', 11)):
    try:
        doc.styles[n].font.size = Pt(sz)
    except Exception:
        pass

CODE_FONT = 'Consolas'

def add_runs(para, text):
    """Parse inline **bold**, *italic*, `code` into runs on `para`."""
    # Tokenize by **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**') and tok.endswith('**'):
            r = para.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = para.add_run(tok[1:-1]); r.font.name = CODE_FONT; r.font.size = Pt(9)
        elif tok.startswith('*') and tok.endswith('*'):
            r = para.add_run(tok[1:-1]); r.italic = True
        else:
            para.add_run(tok)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])

def add_code_block(text, lang=''):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(text)
    r.font.name = CODE_FONT
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1f, 0x2d, 0x3d)
    # Light grey shading on the paragraph
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F4F6F8')
    pPr.append(shd)

def add_table_from_rows(headers, rows):
    n_cols = len(headers)
    if n_cols == 0: return
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = 'Light List Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        para = hdr[i].paragraphs[0]; para.text = ''
        run = para.add_run(h); run.bold = True; run.font.size = Pt(9)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci in range(n_cols):
            cell_text = row[ci] if ci < len(row) else ''
            para = cells[ci].paragraphs[0]; para.text = ''
            add_runs(para, cell_text)
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(9)

def parse_table_row(line):
    # Strip outer pipes, split by `|`, trim cells
    parts = [c.strip() for c in line.strip().strip('|').split('|')]
    return parts

with open(SRC, 'r', encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
N = len(lines)
in_code = False
code_buf = []
code_lang = ''

while i < N:
    raw = lines[i]
    line = raw.rstrip()

    # Fenced code block
    if line.startswith('```'):
        if not in_code:
            in_code = True
            code_lang = line[3:].strip()
            code_buf = []
        else:
            in_code = False
            add_code_block('\n'.join(code_buf), code_lang)
        i += 1
        continue
    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    # Headings
    m = re.match(r'^(#{1,4})\s+(.+)$', line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        h = doc.add_heading('', level=level)
        add_runs(h, text)
        i += 1
        continue

    # Horizontal rule
    if line.strip() in ('---', '***', '___'):
        # Skip as a section divider — add a small spacer paragraph
        doc.add_paragraph()
        i += 1
        continue

    # Tables: detect a pipe-table header followed by a separator row of |---|
    if line.lstrip().startswith('|') and i+1 < N and re.match(r'^\s*\|?[\s:\-|]+\|?\s*$', lines[i+1]) and '---' in lines[i+1]:
        headers = parse_table_row(line)
        i += 2
        rows = []
        while i < N and lines[i].lstrip().startswith('|'):
            rows.append(parse_table_row(lines[i]))
            i += 1
        add_table_from_rows(headers, rows)
        doc.add_paragraph()
        continue

    # Bullets
    m = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
    if m:
        indent = len(m.group(1))
        text = m.group(2).strip()
        para = doc.add_paragraph(style='List Bullet')
        if indent >= 2:
            para.paragraph_format.left_indent = Inches(0.5 + indent*0.1)
        add_runs(para, text)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
    if m:
        text = m.group(2).strip()
        para = doc.add_paragraph(style='List Number')
        add_runs(para, text)
        i += 1
        continue

    # Blank line
    if line.strip() == '':
        i += 1
        continue

    # Regular paragraph
    para = doc.add_paragraph()
    add_runs(para, line)
    i += 1

doc.save(DST)
print(f'wrote {DST}')
